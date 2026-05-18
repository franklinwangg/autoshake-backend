import sys
import json
import os
from pathlib import Path

from dotenv import load_dotenv
from docx import Document
from openai import OpenAI

load_dotenv()
OPENAI_API_KEY = os.getenv("OPENAI_API_KEY")
MODEL = os.getenv("OPENAI_MODEL", "gpt-4o-mini")

if not OPENAI_API_KEY:
    raise RuntimeError("OPENAI_API_KEY not set in .env")

client = OpenAI(api_key=OPENAI_API_KEY)


ROOT = Path(__file__).resolve().parent
MASTER_RESUME_PATH = ROOT / "master_resume.json"
USER_PROFILE_PATH = ROOT / "user_profile.json"


def load_json(path: Path):
    with path.open("r", encoding="utf-8") as f:
        return json.load(f)


def select_and_rewrite_bullets(job_desc: str, job_title: str, company: str, bullets: list):
    """
    Call OpenAI to:
    - pick the most relevant 8–12 bullets
    - rewrite them to match the job language
    """
    system_prompt = (
        "You are an assistant that tailors a software engineer's resume to a specific job. "
        "You are given:\n"
        "1) A list of candidate resume bullets.\n"
        "2) A job title, company name, and job description.\n\n"
        "Pick the 8–12 most relevant bullets and lightly rewrite them to match the job's language "
        "and keywords while keeping them truthful. Do not invent new technologies or experiences.\n"
        "Return only a JSON list of strings under the key 'bullets'."
    )

    bullet_text = "\n".join(f"- {b['bullet']}" for b in bullets)

    user_content = f"""
JOB TITLE: {job_title}
COMPANY: {company}

JOB DESCRIPTION:
{job_desc}

CANDIDATE BULLETS:
{bullet_text}
"""

    resp = client.chat.completions.create(
        model=MODEL,
        messages=[
            {"role": "system", "content": system_prompt},
            {"role": "user", "content": user_content},
        ],
        response_format={"type": "json_object"}
    )

    content = resp.choices[0].message.content
    data = json.loads(content)
    return data.get("bullets", [])


def build_docx(profile: dict, tailored_bullets: list, job_title: str, company: str, out_path: Path):
    doc = Document()

    # Header
    name = profile["name"]
    contact_line = f"{profile['email']} · {profile['phone']} · {profile['location']} · {profile['github']} · {profile['linkedin']}"

    p_name = doc.add_paragraph()
    run = p_name.add_run(name)
    run.bold = True
    run.font.size = run.font.size  # use default, tweak if you want

    doc.add_paragraph(contact_line)

    # Role line
    doc.add_paragraph(f"Target Role: {job_title} @ {company}")

    # Education
    doc.add_heading("Education", level=1)
    for edu in profile.get("education", []):
        p = doc.add_paragraph()
        p.add_run(f"{edu['school']} – {edu['degree']} ({edu['grad_year']})").bold = True
        for d in edu.get("details", []):
            doc.add_paragraph(d, style="List Bullet")

    # Experience / Projects (tailored)
    doc.add_heading("Experience & Projects", level=1)
    for b in tailored_bullets:
        doc.add_paragraph(b, style="List Bullet")

    out_path.parent.mkdir(parents=True, exist_ok=True)
    doc.save(str(out_path))


def main():
    # Read JSON payload from stdin
    raw = sys.stdin.read().strip()
    if not raw:
        print(json.dumps({"error": "No input received"}))
        return

    payload = json.loads(raw)
    job_desc = payload["jobDescription"]
    job_title = payload.get("jobTitle", "Software Engineer")
    company = payload.get("companyName", "Company")
    output_path = Path(payload.get("outputPath", "generated_resume.docx"))

    master = load_json(MASTER_RESUME_PATH)
    profile = load_json(USER_PROFILE_PATH)

    tailored_bullets = select_and_rewrite_bullets(job_desc, job_title, company, master)

    build_docx(profile, tailored_bullets, job_title, company, output_path)

    print(json.dumps({"docxPath": str(output_path.resolve())}))


if __name__ == "__main__":
    main()
